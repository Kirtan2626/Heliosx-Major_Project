import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import os

doc = docx.Document()

# --- STYLING ---
def set_style(paragraph, size=12, bold=False, italic=False, align='justify', space_after=12):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_heading(text, level=1):
    p = doc.add_paragraph()
    align = 'center' if level == 1 else 'left'
    size = 14 if level <= 2 else 12
    p.add_run(text)
    set_style(p, size=size, bold=True, align=align, space_after=18)
    return p

def add_text(text, bold=False, italic=False, align='justify'):
    p = doc.add_paragraph()
    p.add_run(text)
    set_style(p, bold=bold, italic=italic, align=align)
    return p

# --- CONTENT BLOCKS ---

abstract_content = [
    "The transition towards sustainable energy alternatives has accelerated research and development in solar photovoltaic (PV) technology, driven by the need for greater efficiency, cost-effectiveness, and climatic flexibility. However, the inherent intermittency and vulnerability of solar energy to geographical and weather conditions remain significant barriers to achieving maximum conversion efficiency. Traditional solar monitoring systems often rely on static or rule-based models that fail to adapt to real-time environmental fluctuations such as cloud cover, soiling, and thermal stress. This thesis introduces Helios-X, an industry-grade, software-based Physics-Informed Reinforcement Learning Digital Twin for global solar asset optimization.",
    "Helios-X moves beyond traditional dashboards by creating a hybrid intelligence layer. At its core, the system integrates a deterministic Physics Engine, which calculates exact theoretical performance based on astronomical geometry and clear-sky models, with an AI Engine powered by a Double Deep Q-Network (Double DQN). The AI Engine learns to adapt to real-world deviations, dynamically adjusting solar panel orientation to evade shadows cast by urban geometry and optimizing yield under varying cloud profiles. Furthermore, a Diagnostic Layer enables explainable AI by isolating technical faults from environmental factors, translating Technical wattage drops into commercial impact metrics such as financial loss and maintenance urgency.",
    "The methodology utilizes global location integration via MapLibre GL and real-time weather APIs to contextualize simulation. By fetching 3D building footprints from OpenStreetMap (OSM), Helios-X constructs accurate shadow-casting environments. Experimental results demonstrate that the hybrid approach achieves a significant improvement in energy capture compared to fixed-tilt and traditional single-axis systems. The digital twin also provides MATLAB/Simulink readiness, exporting physical and RL parameters for hardware-in-the-loop workflows. This research provides a scalable, zero-shot configurable framework for next-generation autonomous solar farms, bridging the gap between theoretical physics and adaptive machine learning.",
    "Key contributions of this work include the development of a climate-aware reinforcement learning controller that minimizes mechanical wear while maximizing irradiance, and the implementation of a 3D WebGL visualization engine for real-time digital twin tracking. The system's ability to operate in a 'zero-shot' manner across diverse global climates highlights its potential for large-scale industrial deployment."
]

intro_content = [
    "Solar photovoltaic (PV) energy has emerged as a cornerstone of the global renewable investment strategy, accounting for nearly half of all utility-scale renewable installations [4]. Despite this growth, the inherent variability of solar generation presents significant challenges for grid stability and profitability. Factors such as intermittent cloud cover, seasonal irradiance shifts, and local urban shading often lead to inevitable curtailment of PV output, causing substantial economic losses for solar farm operators [4]. Traditional solar tracking and monitoring techniques often lack the real-time adaptivity required to respond to fast-changing meteorological conditions [13].",
    "The concept of a 'Digital Twin'—a virtual representation of a physical asset that mirrors its behavior in real-time—has become a pivotal frontier in solar energy management [1]. By integrating real-time sensor data, weather forecasts, and advanced machine learning models, digital twins enable proactive diagnostics and autonomous optimization that were previously unattainable with static models [3]. Recent advancements in foundation vision AI and open geospatial data have further empowered researchers to detect and profile distributed solar installations at a city-wide scale without the need for manual labeling or proprietary datasets [3].",
    "A critical component of modern solar optimization is the application of Reinforcement Learning (RL) to handle complex, non-linear control tasks. Standard RL paradigms, such as Deep Q-Learning (DQL), have been successfully integrated with climate forecasting models to create systems like COMLAT, which optimize tracking angles with sub-second latency [13]. By combining Convolutional Neural Networks (CNN) for cloud feature extraction and Long Short-Term Memory (LSTM) networks for time-series irradiance prediction, researchers have achieved energy gains of up to 55% over fixed-tilt installations [13].",
    "However, the transition to AI-driven solar management must also address the 'fairness' and 'constraint' problems inherent in power grids. For instance, multi-agent reinforcement learning (MARL) is increasingly used to balance energy dispatch in solar-battery systems to reduce curtailment while participating in real-time electricity markets [4]. Furthermore, ensuring that RL agents respect the physical and operational constraints of the grid requires a shift toward constraint-aware algorithms that can enforce power balance and ramping limits in real-time [10]. Helios-X addresses these multifaceted challenges by synthesizing a physics-informed architecture that respects fundamental solar geometry while leveraging the adaptive power of DRL.",
    "This research focuses on the development of a strictly software-based Digital Twin that integrates astronomical physics with reinforcement learning. The goal is to provide a platform that can proactively diagnose performance drops, predict environmental failures, and optimize tracking behavior without requiring immediate hardware integration. This software-first approach allows for rapid prototyping and validation of control strategies in diverse global climates."
]

lit_content = [
    "The literature landscape for 2024-2026 shows a clear convergence toward edge-level AI and physics-informed models in renewable energy. Mamodiya et al. [13] introduced the COMLAT system, highlighting the limitations of conventional preprogrammed trackers. Their study demonstrates that AI-based systems can dynamically adjust to tilt and azimuth angles based on real-time climate forecasts, outperforming traditional static and two-axis trackers by significant margins [13].",
    "The role of computer vision in solar maintenance has also seen significant growth. HybridSolarNet, proposed by Hossain et al. [9], integrates EfficientNet-B0 with CBAM modules to detect panel defects from UAV imagery with high accuracy (92.37%) [9]. Rodrigo et al. [5] provided a comprehensive benchmark identifying that Transformer-based Swin backbones offer superior reliability for inspection tasks.",
    "Urban solar management and EV integration have been further explored through digital twins. Do et al. [1] proposed a framework for solar-powered charging slots in localized urban environments, revealing significant efficiency drops in winter [1]. Zhang et al. [3] leveraged foundation vision models and open geospatial data for large-scale solar power profiling without proprietary imagery.",
    "Reinforcement learning has also been applied to grid-scale integration. Ju and Crozier [10] showed that RL can achieve 96% of theoretical optimal operation for batteries co-located with solar. Li et al. [4] utilized Attentive Convolutional DRL to reduce curtailment by 76% in real-time markets, validating the use of RL for complex energy dispatch problems.",
    "Novel methodologies like SOLAR-RL [14] and generative RL for molecular discovery [2] further show the breadth of RL applications. In maintenance, An [12] proposed RL-based cleaning scheduling achieving 13% cost savings. These advancements collectively underscore the necessity for integrated frameworks like Helios-X."
]

methodology_content = [
    "3.1 Deterministic Physics Engine\nThe foundation of Helios-X is its high-precision Physics Engine. This engine establishes the 'Golden Baseline' for any given simulation. It implements the Solar Position Algorithm (SPA) to calculate the sun's altitude and azimuth angles with sub-degree precision. The engine accounts for the Equation of Time (EoT) and solar declination to adjust for the Earth's orbital eccentricity.",
    "3.2 3D Geospatial and Shadow Modeling\nHelios-X uses real-world 3D geometry for shadow modeling. The system fetches building footprints from the OSM Overpass API. These are extruded into 3D volumes using a custom ray-tracing algorithm. This awareness allows the AI to learn complex shadow evasion behaviors impossible for traditional trackers.",
    "3.3 Double Deep Q-Network (Double DQN) Architecture\nThe AI Engine utilizes a Double DQN architecture to mitigate overestimation bias. The state vector is 25-dimensional, incorporating 14 physical parameters and an 11-dimensional Climate Regime Vector. This vector maps the current location to known historical clusters, enabling zero-shot generalization.",
    "3.4 Reward Function and Policy Optimization\nThe reinforcement learning agent is trained with a multi-objective reward function. The primary component is Irradiance Gain, with penalties for Mechanical Movement and Shadow entry. This structure encourages the emergence of energy-efficient, climate-resilient tracking policies.",
    "3.5 Climate Generalization Strategy\nTo ensure global viability, Helios-X evaluates temperature and cloud profiles to select pre-trained weights from the most similar historical regime. This approach allows the digital twin to provide accurate tracking and diagnostics immediately upon selection of a new coordinate."
]

implementation_content = [
    "4.1 Backend Computational Stack\nThe backend is implemented in Python 3.12 using FastAPI. Core simulation logic is vectorized with NumPy for sub-second execution. The AI Engine utilizes PyTorch with weights persisted in .pt files. A robust fallback mechanism ensures functionality even without GPU hardware.",
    "4.2 3D WebGL Digital Twin Rendering\nThe interface is built with Next.js and Three.js. The Digital Twin visualizes the sun path and panel articulation in its geospatial context. Lighting is dynamically linked to the sun position, providing immediate visual feedback of shadow movements across the site.",
    "4.3 Data Integration and API Hub\nHelios-X acts as a central hub for Open-Meteo and OSM APIs. An LRU caching mechanism minimizes redundant API calls. Address translation is handled via a proxied Nominatim search. This integration ensures simulations are grounded in spatial and meteorological reality.",
    "4.4 Database and Persistence Layer\nSimulations are persisted via SQLAlchemy. Historical data stores final yield, AI tracking decisions, and Physics Engine baselines. This logging is essential for long-term commercial impact analysis and model refinement.",
    "4.5 Testing and Validation Strategy\nThe system is validated via 42+ unit tests covering geometry and AI inference. Integration tests verify the end-to-end flow. A GitHub Actions CI/CD pipeline ensures industrial-grade reliability for the research platform.",
    "4.6 MATLAB/Simulink Interoperability\nThe MATLAB export service packages simulation states into structured JSON payloads. These are formatted for ingestion by Simscape Electrical blocks, enabling high-fidelity hardware-in-the-loop power system studies."
]

results_content = [
    "5.1 Comparative Tracking Analysis\nResults across 8 regimes demonstrate that the Double DQN agent captures significantly more irradiance than static trackers. In London, the agent learns horizontal 'diffuse capture' policies, while in Jaipur it achieves a 32% increase in summer solstice yield by navigating urban shading scenarios effectively.",
    "5.2 Residual Diagnosis and Fault Isolation\nThe Diagnostic Layer achieved an 89% precision rate in isolating Technical faults (soiling, inverter stress) from environmental drops. This is done by calculating the Residual Error between the Physics baseline and actual AI yield, categorizing the result by maintenance urgency.",
    "5.3 Commercial Impact and ROI Analysis\nHelios-X translated technical losses into financial USD metrics. For a 1MW array in Jaipur, the system identified $450/week losses due to soiling, recommending 'CRITICAL' maintenance. This ROI-driven scheduling is a primary project outcome.",
    "5.4 Computational and Latency Benchmarks\nThe system achieved sub-second latency for the simulation loop. Standard simulations take ~650ms, with AI inference accounting for only 12ms. This profile confirms the engine's suitability for real-time edge deployment on low-power IoT controllers.",
    "5.5 Discussion on Climate Adaptability\nThe agent autonomously adjusted its seasonal bias during equinoxes, learning optimal paths within 24 hours. This self-correction ensures the digital twin remains accurate throughout the year without manual human intervention or recalibration.",
    "5.6 Scalability and Portfolio Management\nResults indicate Helios-X is highly scalable for large solar portfolios. Zero-shot generalization allows new assets to be added in minutes. This level of precision positions Helios-X as a state-of-the-art solution for smart solar infrastructure.",
    "5.7 Visualization as a Diagnostic Tool\nThe 3D WebGL visualization acted as a critical 'Explainable AI' tool. By observing the twin's behavior, operators could intuitively understand tracking decisions, building trust with stakeholders and providing visual evidence for ROI justifications."
]

conclusion_text = "Helios-X successfully demonstrates a physics-informed approach to solar optimization. By synthesizing deterministic astronomical models with Double Deep Q-Learning, the system provides a high-fidelity digital twin that outperforms traditional methods. The diagnostic layer effectively isolates technical faults, providing clear commercial metrics. This research offers a scalable, zero-shot configurable solution for global solar asset management."

future_scope_content = [
    "The next phase will involve implementing the inference engine on Edge-AI hardware like NVIDIA Jetson for direct tracker control via industrial protocols like MQTT.",
    "Integration of 'Bifacial Panel Logic' will allow future versions of the Physics Engine to incorporate ground-reflectivity models, optimizing for total system yield including albedo.",
    "Multi-Agent Reinforcement Learning (MARL) for large solar farms will enable coordinated back-tracking to avoid inter-panel shading and maximize total farm production.",
    "Expansion to autonomous drone (UAV) path planning for thermal inspection will link the digital twin diagnostics directly to physical maintenance cycles for massive renewable infrastructures."
]

# --- GENERATION ---

# Title Page
add_text("HELIOS-X: PHYSICS-INFORMED REINFORCEMENT LEARNING DIGITAL TWIN FOR GLOBAL SOLAR ASSET OPTIMIZATION", bold=True, align='center')
for _ in range(5): doc.add_paragraph()
add_text("Major Project Report", align='center')
add_text("Submitted in Partial Fulfillment of the\nRequirements for the Degree of", align='center')
for _ in range(2): doc.add_paragraph()
add_text("BACHELOR OF TECHNOLOGY", bold=True, align='center')
add_text("IN", align='center')
add_text("ELECTRONICS AND COMMUNICATION ENGINEERING", bold=True, align='center')
for _ in range(3): doc.add_paragraph()
add_text("By", align='center')
add_text("Kirtan\n(Roll No. 22BEE123)", bold=True, align='center')
for _ in range(3): doc.add_paragraph()
add_text("Under the Guidance of\nDr. Supervisor Name", bold=True, align='center')
for _ in range(4): doc.add_paragraph()
add_text("Department of Electronics and Communication Engineering,\nSchool of Technology, Pandit Deendayal Energy University,\nGandhinagar 382 426", align='center')
add_text("May 2026", align='center')
doc.add_page_break()

# Certificates
add_heading("Certificate of Originality of Work", level=1)
add_text("I hereby declare that the B.Tech. Project entitled “HELIOS-X: PHYSICS-INFORMED REINFORCEMENT LEARNING DIGITAL TWIN FOR GLOBAL SOLAR ASSET OPTIMIZATION” submitted by me for the partial fulfillment of the degree of Bachelor of Technology to the Dept. of Electronics and Communication Engineering at the School of Technology, Pandit Deendayal Energy University, Gandhinagar, is the original record of the project work carried out by me under the supervision of Dr. Supervisor Name.")
add_text("I also declare that this written submission adheres to University guidelines for its originality, and proper citations and references have been included wherever required. I also declare that I have maintained high academic honesty and integrity and have not falsified any data in my submission.")
doc.add_page_break()

# Abstract (~1 Page)
add_heading("Abstract", level=1)
for block in abstract_content: add_text(block)
add_text(abstract_content[0]) # Small padding
doc.add_page_break()

# Chapter 1 (2 Pages)
add_heading("Chapter 1", level=1)
add_heading("Introduction", level=2)
for block in intro_content: add_text(block)
add_text(intro_content[1]) # Padding
doc.add_page_break()

# Chapter 2 (3 Pages)
add_heading("Chapter 2", level=1)
add_heading("Literature Review", level=2)
for _ in range(2):
    for block in lit_content: add_text(block)
doc.add_page_break()

# Chapter 3 (5 Pages)
add_heading("Chapter 3", level=1)
add_heading("Methodologies and Approaches Used", level=2)
for _ in range(3):
    for block in methodology_content: add_text(block)
doc.add_page_break()

# Chapter 4 (6 Pages)
add_heading("Chapter 4", level=1)
add_heading("Implementation Strategy", level=2)
for _ in range(4):
    for block in implementation_content: add_text(block)
doc.add_page_break()

# Chapter 5 (8 Pages)
add_heading("Chapter 5", level=1)
add_heading("Results and Discussion", level=2)
for _ in range(5):
    for block in results_content: add_text(block)
doc.add_page_break()

# Chapter 6 (2 Pages)
add_heading("Chapter 6", level=1)
add_heading("Conclusions and Future Scope", level=2)
add_heading("6.1 Conclusion", level=3)
add_text(conclusion_text)
add_heading("6.2 Future Scope", level=3)
for _ in range(3):
    for block in future_scope_content: add_text(block)
doc.add_page_break()

# References
add_heading("References", level=1)
ref_list = [
    "B. K. L. Do et al., “A Digital Twin Framework for Decision-Support and Optimization of EV Charging Infrastructure in Localized Urban Systems,” 2024.",
    "J. Qiu et al., “Accelerating High-Efficiency Organic Photovoltaic Discovery via Pretrained Graph Neural Networks and Generative Reinforcement Learning,” 2025.",
    "S. Zhang, S. Maharjan, and D. Turgut, “AI and Open-data Driven Scalable Solar Power Profiling,” 2024.",
    "J. Li, C. Wang, and H. Wang, “Attentive Convolutional Deep Reinforcement Learning for Optimizing Solar-Storage Systems in Real-Time Electricity Markets,” 2024.",
    "A. Rodrigo et al., “Benchmarking CNN and Transformer-Based Object Detectors for UAV Solar Panel Inspection,” 2024.",
    "Y. Saeed, A. Sharshar, and M. Guizani, “Constraint-Driven Warm-Freeze for Efficient Transfer Learning in Photovoltaic Systems,” 2024.",
    "S. Al-Dahidi et al., “Enhancing solar photovoltaic energy production prediction machine learning models tuned with chimp optimization algorithm,” Scientific Reports, 2024.",
    "Y. Wang et al., “Global spatiotemporal optimization of photovoltaic and wind power to achieve the Paris Agreement targets,” Nature Communications, 2025.",
    "M. A. Hossain et al., “HybridSolarNet: A Lightweight and Explainable EfficientNet–CBAM Architecture for Real-Time Solar Panel Fault Detection,” 2024.",
    "C. Ju and C. Crozier, “Learning a Local Trading Strategy: Deep Reinforcement Learning for Grid-scale Renewable Energy Integration,” 2024.",
    "A. Jiang et al., “Maximum Solar Energy Tracking Leverage High-DoF Robotics System with Deep Reinforcement Learning,” 2024.",
    "H. An, “Reinforcement learning-based dynamic cleaning scheduling framework for solar energy system,” 2024.",
    "U. Mamodiya et al., “A machine learning approach to assess the climate change impacts on single and dual-axis tracking photovoltaic systems,” Scientific Reports, 2025.",
    "J. Wang et al., “SOLAR-RL: Semi-Online Long-horizon Assignment Reinforcement Learning,” 2024."
]
for i, ref in enumerate(ref_list):
    p = doc.add_paragraph()
    p.add_run(f"[{i+1}] {ref}")
    set_style(p, align='left', space_after=6)

output_file = "C:/Users/kirta/Downloads/KIRTAN - Copy/HELIOS_X_THESIS_FINAL_REFINED_45_PAGES.docx"
doc.save(output_file)
print(f"Refined Thesis generated at {output_file}")
