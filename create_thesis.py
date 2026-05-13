import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

doc = docx.Document()

# Styles
style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'Times New Roman'
font.size = Pt(12)

style_heading = doc.styles['Heading 1']
font_heading = style_heading.font
font_heading.name = 'Times New Roman'
font_heading.size = Pt(14)
font_heading.bold = True

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    return p

def add_paragraph(text, align='justify', bold=False):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if bold:
        run.bold = True
    return p

# TITLE PAGE
add_title("HELIOS-X: SOFTWARE DIGITAL TWIN FOR SOLAR ASSET DIAGNOSIS AND OPTIMIZATION")
for _ in range(2): doc.add_paragraph()
add_paragraph("Major Project/Comprehensive Project Report", align='center')
add_paragraph("Submitted in Partial Fulfillment of the\nRequirements for the Degree of", align='center')
for _ in range(2): doc.add_paragraph()
add_paragraph("BACHELOR OF TECHNOLOGY", align='center', bold=True)
add_paragraph("IN", align='center')
add_paragraph("ELECTRONICS AND COMMUNICATION ENGINEERING", align='center', bold=True)
for _ in range(2): doc.add_paragraph()
add_paragraph("By", align='center')
add_paragraph("Kirtan\n(Roll No. 22BEE123)", align='center', bold=True)
for _ in range(2): doc.add_paragraph()
add_paragraph("Under the Guidance of\nDr. Supervisor Name", align='center', bold=True)
for _ in range(4): doc.add_paragraph()
add_paragraph("Department of Electronics and Communication Engineering,\nSchool of Technology, Pandit Deendayal Energy University,\nGandhinagar 382 426", align='center')
for _ in range(2): doc.add_paragraph()
add_paragraph("May 2026", align='center')
doc.add_page_break()

# CERTIFICATE OF ORIGINALITY
add_title("Certificate of Originality of Work")
doc.add_paragraph()
add_paragraph("I hereby declare that the B.Tech. Project entitled “Helios-X: Software Digital Twin for Solar Asset Diagnosis and Optimization” submitted by me for the partial fulfillment of the degree of Bachelor of Technology to the Dept. of Electronics and Communication Engineering at the School of Technology, Pandit Deendayal Energy University, Gandhinagar, is the original record of the project work carried out by me under the supervision of Dr. Supervisor Name.")
doc.add_paragraph()
add_paragraph("I also declare that this written submission adheres to University guidelines for its originality, and proper citations and references have been included wherever required.")
doc.add_paragraph()
add_paragraph("I also declare that I have maintained high academic honesty and integrity and have not falsified any data in my submission.")
doc.add_paragraph()
add_paragraph("I also understand that violation of any guidelines in this regard will attract disciplinary action by the institute.")
for _ in range(2): doc.add_paragraph()
add_paragraph("Name of the Student: Kirtan")
add_paragraph("Roll Number of the Student: 22BEE123")
add_paragraph("Signature of the Student:")
add_paragraph("Name of the Supervisor: Dr. Supervisor Name")
add_paragraph("Designation of the Supervisor: Assistant Professor")
add_paragraph("Signature of the Supervisor:")
add_paragraph("Place: Gandhinagar")
add_paragraph("Date: 10 May 2026")
doc.add_page_break()

# CERTIFICATE FROM SUPERVISOR
add_title("Certificate from the Project Supervisor/Head")
doc.add_paragraph()
add_paragraph("This is to certify that the Major/Comprehensive Project Report entitled “Helios-X: Software Digital Twin for Solar Asset Diagnosis and Optimization” submitted by Mr. Kirtan, Roll No. 22BEE123 towards the partial fulfilment of the requirements for the award of degree in Bachelor of Technology in the field of Electronics and Communication Engineering from the School of Technology, Pandit Deendayal Energy University, Gandhinagar is the record of work carried out by him under my supervision and guidance. The work submitted by the student has in my opinion reached a level required for being accepted for examination. The results embodied in this major project work to the best of our knowledge have not been submitted to any other University or Institution for the award of any degree or diploma.")
doc.add_page_break()

# ACKNOWLEDGEMENT
add_title("Acknowledgement")
doc.add_paragraph()
add_paragraph("I would like to express my sincere gratitude to my respected supervisor Dr. Supervisor Name for their invaluable guidance, constant encouragement, and insightful suggestions throughout the course of this B. Tech. thesis. Their support and expertise played a crucial role in shaping this work. I am also deeply thankful to the examiner for their careful evaluation, constructive feedback, and thoughtful recommendations, which have greatly contributed to improving the quality of this thesis. Their time and effort are truly appreciated.")
for _ in range(2): doc.add_paragraph()
add_paragraph("(Kirtan)", align='right')
doc.add_page_break()

# ABSTRACT
add_title("Abstract")
doc.add_paragraph()
add_paragraph("Large commercial solar installations often suffer from yield degradation due to complex, overlapping factors: thermal stress, dust accumulation (soiling), partial shading from urban geometry, and inverter or panel faults. This project, Helios-X, aims to solve this by creating a Physics-Informed Reinforcement Learning Digital Twin for solar energy systems. It moves beyond traditional, static solar dashboards by proactively diagnosing performance drops, predicting environmental failures, optimizing panel tracking to evade shadows, and providing commercial impact analytics—all without requiring immediate physical hardware integration.")
doc.add_paragraph()
add_paragraph("Helios-X creates a hybrid intelligence layer composed of a Physics Engine that calculates exact theoretical performance based on clear-sky conditions and astronomical geometry, an AI Engine (Double Deep Q-Network) that adapts to real-world deviations to dynamically adjust panel orientation, and a Diagnostic Layer that compares theoretical outputs to AI simulated reality to isolate and explain the causes of energy loss. The platform enables users to select any location globally, fetches real-time data, and visually represents the tracking behavior and shadow-casting using a 3D WebGL visualization. Furthermore, it quantifies technical issues into commercial impact metrics such as financial loss and maintenance urgency.")
doc.add_page_break()

# LISTS
add_title("INDEX")
add_paragraph("(Auto-generated TOC to be inserted here in word processor)")
doc.add_page_break()

add_title("LIST OF FIGURES")
add_paragraph("Figure 1.1: Helios-X Architecture... (Page 10)")
doc.add_page_break()

add_title("LIST OF TABLES")
add_paragraph("Table 3.1: Physics Engine Parameters... (Page 20)")
doc.add_page_break()

add_title("NOMENCLATURE")
add_paragraph("DQN - Deep Q-Network\nRL - Reinforcement Learning\nOSM - OpenStreetMap\nAPI - Application Programming Interface")
doc.add_page_break()

# CHAPTER 1
doc.add_heading('Chapter 1', level=1)
doc.add_heading('Introduction', level=2)
doc.add_heading('1.1 Prologue', level=3)
add_paragraph("The explosive growth of smart grids and digital twin technologies represents a revolutionary change in the management of renewable energy resources [1]. Modern commercial solar installations are increasingly complex and vulnerable to various degradation factors such as shading, soiling, and thermal stress. The need for proactive diagnostics and automated optimization has driven the development of software-based simulation models [2].")
doc.add_heading('1.2 Project Goal', level=3)
add_paragraph("The primary goal of Helios-X is to provide an industry-grade, strictly software-based Physics-Informed Reinforcement Learning Digital Twin for solar energy systems. It proactively diagnoses performance drops, predicts failures, optimizes panel tracking, and provides commercial impact analytics.")
doc.add_heading('1.3 Scope', level=3)
add_paragraph("The scope of Helios-X is bound strictly to the Software Simulation and Visualization Layer. It is designed to be a defensible, presentation-ready platform suitable for patent discussions, hackathons, and commercial investor demonstrations. Features like direct hardware control and live SCADA data ingestion from physical inverters are currently out of scope and reserved for future expansion.")
doc.add_page_break()

# CHAPTER 2
doc.add_heading('Chapter 2', level=1)
doc.add_heading('Literature Review', level=2)
add_paragraph("Recent studies in renewable energy emphasize the integration of Deep Reinforcement Learning (DRL) for system optimization. Data-driven models and digital twins have successfully managed to forecast solar irradiance, detect faults visually using convolutional networks, and optimize operations under partial shading conditions [3].")
add_paragraph("Physics-informed AI has been established as a crucial step toward accurate solar degradation modeling. By forcing deep learning models to adhere to physical laws—such as astronomical geometry and irradiance equations—generalization in unseen climates is vastly improved [4].")
doc.add_page_break()

# CHAPTER 3
doc.add_heading('Chapter 3', level=1)
doc.add_heading('Methodology and Architecture', level=2)
doc.add_heading('3.1 The Hybrid Intelligence Layer', level=3)
add_paragraph("Helios-X employs a three-tiered architecture:")
add_paragraph("1. The Physics Engine: Calculates theoretical clear-sky yield and precise solar geometry.")
add_paragraph("2. The AI Engine: Uses a Double Deep Q-Network (Double DQN) PyTorch model to learn optimal panel tracking strategies in environments constrained by urban geometry and cloud cover.")
add_paragraph("3. The Diagnostic Layer: Isolates faults by calculating the residual between the deterministic baseline and the AI's actual simulated yield.")
doc.add_heading('3.2 Real-Time Data Integration', level=3)
add_paragraph("The backend fetches real-world temperature, humidity, wind, and cloud cover from Open-Meteo and OpenWeatherMap APIs. Additionally, OpenStreetMap (OSM) Overpass API is used to fetch building footprints and tree data to construct accurate 3D shadow-casting geometry.")
doc.add_page_break()

# CHAPTER 4
doc.add_heading('Chapter 4', level=1)
doc.add_heading('Implementation Details', level=2)
doc.add_heading('4.1 Technology Stack', level=3)
add_paragraph("The frontend is built using Next.js with TypeScript and TailwindCSS, incorporating MapLibre GL for geospatial mapping and Three.js for 3D Digital Twin rendering. The backend uses Python 3 for handling lightweight HTTP requests, running the pure-Python Physics Engine, and serving the PyTorch-based AI models.")
doc.add_heading('4.2 3D Geometry and Visualization', level=3)
add_paragraph("The 3D viewer traces exact polygon footprints returned by the OSM API using THREE.ExtrudeGeometry, rendering accurate real-world building shapes. The visualization also handles nighttime rendering smoothly with ambient lighting adjustments and provides clear indicators for the sun's position and solar panel tracking behavior.")
doc.add_page_break()

# CHAPTER 5
doc.add_heading('Chapter 5', level=1)
doc.add_heading('Result and Discussions', level=2)
doc.add_heading('5.1 Commercial Impact Analytics', level=3)
add_paragraph("Helios-X successfully translates raw physics wattage drops into actionable business metrics: Estimated Daily kWh Loss, Financial USD Loss, and Maintenance Urgency. These metrics are dynamically visualized in the Next.js frontend, providing investors and stakeholders with immediate clarity regarding asset performance.")
doc.add_heading('5.2 Fault Diagnosis', level=3)
add_paragraph("The diagnostic engine consistently separates environmental phenomena (like passing clouds) from physical defects (like soiling and thermal stress). The integration of Python-based models with asynchronous SQLite/Postgres persistence ensures all simulations are logged successfully.")
doc.add_page_break()

# CHAPTER 6
doc.add_heading('Chapter 6', level=1)
doc.add_heading('Conclusion', level=2)
add_paragraph("Helios-X presents a state-of-the-art solution for solar asset management. By synthesizing a deterministic Physics Engine with advanced Reinforcement Learning and robust 3D geospatial rendering, it creates a high-fidelity Software Digital Twin. It effectively addresses the complex challenges of urban shading, hardware fault diagnosis, and financial impact estimation in a presentation-ready, zero-shot configurable framework.")
doc.add_page_break()

# CHAPTER 7
doc.add_heading('Chapter 7', level=1)
doc.add_heading('Future Prospects', level=2)
add_paragraph("Future development will focus on scaling the application to handle direct hardware control via IoT actuator telemetry. On the software side, moving frontend state management to React Context or Zustand, introducing user authentication, and deploying the system to AWS ECS or a VPS are prioritized. Real-time SCADA data ingestion from physical inverters will bridge the final gap between the software simulation layer and active hardware.")
doc.add_page_break()

# REFERENCES
add_title("References")
add_paragraph("[1] J. K. Author, “Title of chapter in the book,” in Title of His Published Book, 2nd ed. City: Publisher, 2025, ch.1, sec. 2, pp. 10–15.")
add_paragraph("[2] A. Researcher, “Digital Twins in Solar Systems,” J. Renew. Energy, vol. 14, no. 2, pp. 100-110, Feb. 2025.")
add_paragraph("[3] S. Scholar, “Reinforcement learning for smart grids,” in Proc. IEEE Power Conf., New York, NY, 2024, pp. 40-45.")
add_paragraph("[4] M. Scientist, “Physics-informed neural networks in photovoltaic systems,” IEEE Trans. Smart Grid, vol. 16, no. 1, pp. 50-60, Jan. 2026.")
doc.add_page_break()

# APPENDIX
add_title("Appendix")
doc.add_paragraph()
add_paragraph("A. MATLAB/Simulink Integration Payload JSON Format")
add_paragraph("The system exports simulation states into a structured JSON payload ready for ingestion by MATLAB Simscape Electrical workflows.")

doc.save("C:/Users/kirta/Downloads/KIRTAN - Copy/CP_MP_Thesis_format_final_2026_Heliox.docx")
print("Thesis generated successfully.")
